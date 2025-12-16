import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import threading
import pandas as pd
import os
import json
import re
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from openai import OpenAI

# ================= Configuration =================
class Config:
    SPLIT_ANCHOR = "RECIPIENT FULL ADDRESS" 
    COL_IDX_ADDRESS = 3  # Column D (Index 3)
    COL_IDX_MESSAGE = 8  # Column I (Index 8)
    
    IGNORE_KEYWORDS = [
        "form instructions",
        "basket name",
        "delivery date",
        "special instructions",
        "must be a valid address"
    ]
    
    USER_BLACKLIST = [
        "750 millway", 
        "my baskets", 
        "unit #4"
    ]

# ================= Logic Class =================
class CardGenerator:
    def __init__(self, log_callback, progress_callback):
        self.log = log_callback
        self.progress_cb = progress_callback
        self.api_key = None
        self.base_url = None
        self.model_name = "gpt-3.5-turbo"

    def update_settings(self, api_key, base_url, model_name):
        self.api_key = api_key
        self.base_url = base_url if base_url.strip() else "https://api.openai.com/v1"
        self.model_name = model_name if model_name.strip() else "gpt-3.5-turbo"

    def _is_garbage(self, text):
        t = text.lower()
        if not t.strip(): return True
        for kw in Config.IGNORE_KEYWORDS:
            if kw in t: return True
        return False

    def _clean_labels(self, text):
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            new_line = re.sub(r'(?i)^\s*(tel|phone|ph|mobile|cell)[:\.]?\s*', '', line)
            if new_line.strip() and new_line.strip() not in [",", ".", "-"]:
                cleaned_lines.append(new_line)
        return "\n".join(cleaned_lines)

    def _is_ai_chatting(self, response_text):
        bad_starts = ["sure!", "sure,", "here is", "please provide", "i cannot", "sorry", "certainly"]
        lower_resp = response_text.lower().strip()
        for bad in bad_starts:
            if lower_resp.startswith(bad):
                return True
        return False

    # --- AI: Format Address ---
    def ai_format_block(self, raw_text):
        if not self.api_key or len(raw_text.strip()) < 5: return raw_text
        try:
            client = OpenAI(api_key=self.api_key, base_url=self.base_url)
            prompt = f"""
            Format this shipping address block.
            Input:
            ---
            {raw_text}
            ---
            Rules:
            1. STRICTLY REMOVE labels like "TEL:", "Attention:". Keep only values.
            2. Do not write "[insert number]".
            3. Put phone number on the last line.
            4. Output plain text only. 
            5. IF INPUT IS EMPTY OR NONSENSE, RETURN EMPTY STRING.
            """
            response = client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1
            )
            result = response.choices[0].message.content.strip()
            if self._is_ai_chatting(result): return raw_text
            return result
        except Exception as e:
            self.log(f"AI Address Error: {e}")
            return raw_text

    # --- AI: Format Message ---
    def ai_format_message(self, raw_msg):
        if not self.api_key or len(raw_msg.strip()) < 2: return raw_msg
        try:
            client = OpenAI(api_key=self.api_key, base_url=self.base_url)
            prompt = f"""
            Refine this Gift Card Message which was split across multiple Excel rows.
            Input:
            ---
            {raw_msg}
            ---
            Rules:
            1. Merge broken sentences into a single coherent paragraph.
            2. Keep the "From: [Name]" or "Love, [Name]" on a separate line at the bottom.
            3. Fix spacing and punctuation.
            4. Do NOT add quotes or conversational text.
            """
            response = client.chat.completions.create(
                model=self.model_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            result = response.choices[0].message.content.strip()
            if self._is_ai_chatting(result): return raw_msg
            return result
        except Exception as e:
            self.log(f"AI Message Error: {e}")
            return raw_msg

    # ================= STRICT GRID LOGIC (V15) =================
    def read_excel_strict(self, file_path):
        self.log(f"Reading file (Strict Grid Mode): {os.path.basename(file_path)}...")
        self.progress_cb(0)
        
        try:
            df = pd.read_excel(file_path, header=None, dtype=str)
            total_rows = len(df)
            
            # Find First Anchor (Start Point)
            start_idx = 0
            found = False
            for r in range(len(df)):
                val = str(df.iloc[r, Config.COL_IDX_ADDRESS]).strip()
                if Config.SPLIT_ANCHOR.lower() in val.lower():
                    start_idx = r
                    found = True
                    break
            
            if not found:
                self.log("❌ Could not find 'RECIPIENT FULL ADDRESS' header.")
                return []

            parsed_cards = []
            
            # Start logic:
            # If we found header at Row X, the first person starts at X+1.
            current_row = start_idx 
            
            while current_row < total_rows:
                # Update Progress
                prog_val = (current_row / total_rows) * 100
                self.progress_cb(prog_val)

                # --- 1. Header Check (Reset Point) ---
                # Check if the CURRENT row is a header
                cell_val = str(df.iloc[current_row, Config.COL_IDX_ADDRESS]).strip()
                if Config.SPLIT_ANCHOR.lower() in cell_val.lower():
                    self.log(f"🔄 Header found at Row {current_row+1}. Resetting alignment.")
                    current_row += 1  # Move to the row immediately after header (D65 -> D66)
                    continue

                # --- 2. Safety: Look Ahead inside the block ---
                # Before we blindly grab 5 rows, make sure a Header isn't lurking inside them.
                # Example: We only have 3 rows of data, then a new header appears.
                rows_to_check = 5
                if current_row + 5 > total_rows:
                    rows_to_check = total_rows - current_row

                header_interruption_offset = -1
                for offset in range(rows_to_check):
                    r_check = current_row + offset
                    val_check = str(df.iloc[r_check, Config.COL_IDX_ADDRESS]).strip()
                    if Config.SPLIT_ANCHOR.lower() in val_check.lower():
                        header_interruption_offset = offset
                        break
                
                if header_interruption_offset != -1:
                    # A header cut us off! Jump directly to that header and restart loop.
                    # This handles incomplete blocks (e.g., only 3 lines of data then header).
                    self.log(f"⚠️ Block interrupted by header at Row {current_row + header_interruption_offset + 1}.")
                    current_row += header_interruption_offset
                    continue

                # --- 3. Strict Extraction ---
                # If we are here, the next 5 rows are safe (or until EOF).
                block_addr_lines = []
                block_msg_lines = []
                is_block_empty = True

                for k in range(5):
                    r_idx = current_row + k
                    if r_idx >= total_rows: break # EOF safety

                    raw_a = str(df.iloc[r_idx, Config.COL_IDX_ADDRESS]).strip()
                    if raw_a.lower() == 'nan': raw_a = ""
                    
                    raw_m = str(df.iloc[r_idx, Config.COL_IDX_MESSAGE]).strip()
                    if raw_m.lower() == 'nan': raw_m = ""

                    if raw_a or raw_m: is_block_empty = False
                    if raw_a: block_addr_lines.append(raw_a)
                    if raw_m: block_msg_lines.append(raw_m)

                # If the block has data, process it.
                if not is_block_empty:
                    full_addr = "\n".join(block_addr_lines)
                    full_msg = "\n".join(block_msg_lines)
                    # Simple validation: Address needs to be substantial
                    if self._validate_block(full_addr):
                        self._add_card_to_list(parsed_cards, full_addr, full_msg)

                # --- 4. Move Forward Strictly ---
                # Regardless of whether it was empty or full, we strictly move 5 rows.
                current_row += 5

            self.progress_cb(100)
            self.log(f"📊 Extraction Complete. Valid Cards: {len(parsed_cards)}")
            return parsed_cards

        except Exception as e:
            self.log(f"❌ Error Reading File: {e}")
            import traceback
            traceback.print_exc()
            return []

    def _validate_block(self, addr_text):
        if len(addr_text) < 5: return False 
        if self._is_garbage(addr_text): return False
        # Must contain digits (zip code or phone) to be a valid address block
        if not any(char.isdigit() for char in addr_text): return False
        return True

    def _add_card_to_list(self, card_list, raw_addr, raw_msg):
        clean_addr = self._clean_labels(raw_addr)
        for bw in Config.USER_BLACKLIST:
            if bw in clean_addr.lower(): return

        final_addr = clean_addr
        final_msg = raw_msg

        if self.api_key:
            if len(clean_addr) > 5:
                ai_a = self.ai_format_block(clean_addr)
                if ai_a: final_addr = ai_a
                final_addr = self._clean_labels(final_addr)

            if len(raw_msg.strip()) > 2:
                ai_m = self.ai_format_message(raw_msg)
                if ai_m: final_msg = ai_m

        card_list.append({'address': final_addr, 'message': final_msg})

    def generate_word(self, data_list, output_path):
        if not data_list: return

        self.log("Generating Word Document...")
        doc = docx.Document()
        section = doc.sections[0]
        section.page_height = Cm(29.7); section.page_width = Cm(21.0)
        
        margin = Cm(1.0)
        section.top_margin = margin; section.bottom_margin = margin
        section.left_margin = margin; section.right_margin = margin

        ROW_HEIGHT = Cm(6.75)
        COL_WIDTH = Cm(9.5)

        chunk_size = 4
        total_pages = (len(data_list) + chunk_size - 1) // chunk_size

        for page_idx in range(total_pages):
            start = page_idx * chunk_size
            end = start + chunk_size
            page_data = data_list[start:end]

            table = doc.add_table(rows=4, cols=2)
            table.autofit = False 
            
            for row in table.rows:
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = ROW_HEIGHT
                row.cells[0].width = COL_WIDTH
                row.cells[1].width = COL_WIDTH

            for i, item in enumerate(page_data):
                row = table.rows[i]
                
                # Address
                cell_l = row.cells[0]
                cell_l.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p_l = cell_l.paragraphs[0]
                p_l.text = item['address']
                p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if p_l.runs:
                    run_l = p_l.runs[0]
                    run_l.font.name = 'Arial' 
                    run_l.font.size = Pt(10.5) 
                    if i == 3: p_l.paragraph_format.space_before = Pt(28)
                
                # Message
                cell_r = row.cells[1]
                cell_r.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p_r = cell_r.paragraphs[0]
                p_r.text = item['message']
                p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if i == 3: p_r.paragraph_format.space_before = Pt(28)

                if p_r.runs:
                    run_r = p_r.runs[0]
                    run_r.font.name = 'Arial'
                    length = len(item['message'])
                    if length < 30: run_r.font.size = Pt(16)
                    elif length < 80: run_r.font.size = Pt(13)
                    else: run_r.font.size = Pt(10)

            if page_idx < total_pages - 1: doc.add_page_break()
        
        try:
            doc.save(output_path)
            self.log(f"✅ Saved Successfully: {output_path}")
            messagebox.showinfo("Success", "Processing Complete!")
        except Exception as e:
            self.log(f"❌ Error Saving File: {e}")

# ================= UI Class =================
class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Card Generator (V15 - Strict Grid Mode)")
        self.root.geometry("650x700")
        
        self.api_key = tk.StringVar()
        self.base_url = tk.StringVar()
        self.model_name = tk.StringVar(value="gpt-3.5-turbo") 
        self.file_path = tk.StringVar()
        
        self.logic = CardGenerator(self.log_msg, self.update_progress)
        self._load_config()
        self._setup_ui()

    def _load_config(self):
        if os.path.exists("config_ai.json"):
            try:
                with open("config_ai.json", "r") as f:
                    c = json.load(f)
                    self.api_key.set(c.get("api_key", ""))
                    self.base_url.set(c.get("base_url", ""))
                    self.model_name.set(c.get("model_name", "gpt-3.5-turbo"))
            except: pass

    def _save_config(self):
        with open("config_ai.json", "w") as f:
            json.dump({
                "api_key": self.api_key.get(),
                "base_url": self.base_url.get(),
                "model_name": self.model_name.get()
            }, f)

    def _setup_ui(self):
        f_ai = tk.LabelFrame(self.root, text="AI Settings", padx=10, pady=10)
        f_ai.pack(fill="x", padx=10, pady=5)
        
        tk.Label(f_ai, text="API Key:").grid(row=0, column=0, sticky="e")
        tk.Entry(f_ai, textvariable=self.api_key, width=40, show="*").grid(row=0, column=1, padx=5, pady=2)
        tk.Label(f_ai, text="Base URL:").grid(row=1, column=0, sticky="e")
        tk.Entry(f_ai, textvariable=self.base_url, width=40).grid(row=1, column=1, padx=5, pady=2)
        tk.Label(f_ai, text="Model:").grid(row=2, column=0, sticky="e")
        tk.Entry(f_ai, textvariable=self.model_name, width=40).grid(row=2, column=1, padx=5, pady=2)

        f_file = tk.LabelFrame(self.root, text="File Selection", padx=10, pady=10)
        f_file.pack(fill="x", padx=10, pady=5)
        tk.Entry(f_file, textvariable=self.file_path, width=50).pack(side="left")
        tk.Button(f_file, text="Browse Excel", command=self.sel_file).pack(side="left", padx=5)

        f_prog = tk.Frame(self.root, padx=10, pady=5)
        f_prog.pack(fill="x")
        tk.Label(f_prog, text="Processing Progress:").pack(anchor="w")
        self.progress_bar = ttk.Progressbar(f_prog, orient="horizontal", length=100, mode="determinate")
        self.progress_bar.pack(fill="x", pady=2)

        tk.Button(self.root, text="🚀 Start Processing", command=self.run_thread, 
                  bg="#28a745", fg="white", font=("Arial", 12, "bold"), height=2).pack(fill="x", padx=20, pady=10)

        self.log_area = scrolledtext.ScrolledText(self.root, height=15)
        self.log_area.pack(fill="both", expand=True, padx=10, pady=5)

    def sel_file(self):
        p = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx *.xls")])
        if p: self.file_path.set(p)

    def log_msg(self, msg):
        self.root.after(0, lambda: self.log_area.insert(tk.END, f">> {msg}\n"))
        self.root.after(0, lambda: self.log_area.see(tk.END))

    def update_progress(self, val):
        self.root.after(0, lambda: self.progress_bar.configure(value=val))

    def run_thread(self):
        self._save_config()
        self.progress_bar['value'] = 0 
        threading.Thread(target=self.run, daemon=True).start()

    def run(self):
        f = self.file_path.get()
        if not f: return
        self.logic.update_settings(self.api_key.get(), self.base_url.get(), self.model_name.get())
        data = self.logic.read_excel_strict(f)
        if data:
            base = os.path.splitext(f)[0]
            out = f"{base}_English_V15_Strict.docx"
            self.logic.generate_word(data, out)

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()
